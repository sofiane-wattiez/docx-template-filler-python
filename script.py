from docx import Document
import tkinter as tk
from tkinter import messagebox
import os  # Pour gérer les répertoires
import datetime  # Pour récupérer la date actuelle

def replace_text_in_doc(template_path, output_path, replacements):
    # Créer le répertoire de sortie si nécessaire
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document(template_path)
    
    # Remplacer le texte dans les paragraphes
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
    
    # Remplacer le texte dans les cellules des tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    # Sauvegarder le document modifié
    doc.save(output_path)

def generate_documents(data):
    remise_template = "template_remise.docx"
    restitution_template = "template_restitution.docx"
    
    # Vérification de l'existence des fichiers modèle
    if not os.path.exists(remise_template):
        print(f"Le fichier modèle {remise_template} est manquant.")
        messagebox.showerror("Erreur", f"Le fichier modèle {remise_template} est manquant.")
        return

    if not os.path.exists(restitution_template):
        print(f"Le fichier modèle {restitution_template} est manquant.")
        messagebox.showerror("Erreur", f"Le fichier modèle {restitution_template} est manquant.")
        return
    
    # Créer le nom du dossier principal basé sur le nom et prénom
    main_folder_name = f"{data['nom']}_{data['prenom']}"
    output_dir = os.path.join(os.getcwd(), main_folder_name)  # Répertoire principal basé sur le nom et prénom
    
    # Créer le dossier principal si nécessaire
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Nom du fichier de sortie pour le document de remise
    remise_output = os.path.join(main_folder_name, f"Document_Remise_MI_{data['nom']}_{data['prenom']}_{data['new_numero_identification']}_{data['new_numero_de_serie']}.docx")
    
    remise_replacements = {
        "{nom}": data["nom"],
        "{prenom}": data["prenom"],
        "{date}": data["date"],  # La date actuelle
        "{marque}": data["marque"],
        "{modele}": data["modele"],
        "{new_numero_identification}": data["new_numero_identification"],
        "{new_numero_de_serie}": data["new_numero_de_serie"],
    }
    
    try:
        replace_text_in_doc(remise_template, remise_output, remise_replacements)
        print(f"Document généré : {remise_output}")
    except Exception as e:
        print(f"Erreur lors de la génération du document de remise : {e}")
        messagebox.showerror("Erreur", "Erreur lors de la génération du document de remise.")
    
    # Document de restitution, si applicable
    if all(value is not None for value in [data["old_numero_identification"], data["old_numero_de_serie"], data["old_marque"], data["old_modele"]]):
        restitution_output = os.path.join(main_folder_name, f"Document_Restitution_MI_{data['nom']}_{data['prenom']}_{data['old_numero_identification']}_{data['old_numero_de_serie']}.docx")
        restitution_replacements = {
            "{nom}": data["nom"],
            "{prenom}": data["prenom"],
            "{date}": data["date"],  # La date actuelle
            "{old_marque}": data["old_marque"],  # Ancienne marque
            "{old_modele}": data["old_modele"],  # Ancien modèle
            "{old_numero_identification}": data["old_numero_identification"],
            "{old_numero_de_serie}": data["old_numero_de_serie"],
        }
        
        try:
            replace_text_in_doc(restitution_template, restitution_output, restitution_replacements)
            print(f"Document généré : {restitution_output}")
        except Exception as e:
            print(f"Erreur lors de la génération du document de restitution : {e}")
            messagebox.showerror("Erreur", "Erreur lors de la génération du document de restitution.")
    else:
        print("Aucun ancien poste détecté, le document de restitution ne sera pas généré.")

def submit_form():
    # Récupérer les valeurs des champs et les organiser dans un dictionnaire
    data = {
        "nom": entry_nom.get(),
        "prenom": entry_prenom.get(),
        "date": entry_date.get(),  # La date sera récupérée automatiquement ici
        "marque": entry_marque.get(),
        "modele": entry_modele.get(),
        "new_numero_identification": entry_new_id.get(),
        "new_numero_de_serie": entry_new_sn.get(),
        "old_numero_identification": entry_old_id.get() if entry_old_id.get().strip() != "" else None,
        "old_numero_de_serie": entry_old_sn.get() if entry_old_sn.get().strip() != "" else None,
        "old_marque": entry_old_marque.get() if entry_old_marque.get().strip() != "" else None,  # Ancienne marque
        "old_modele": entry_old_modele.get() if entry_old_modele.get().strip() != "" else None,  # Ancien modèle
    }
    generate_documents(data)
    messagebox.showinfo("Succès", "Documents générés avec succès !")

# Interface graphique
root = tk.Tk()
root.title("Générateur de documents")

# Valeurs par défaut pour les champs
default_values = {
    "Nom": "John",
    "Prénom": "Doe",
    "Date": "",  # La date sera remplie automatiquement
    "Marque": "HP",
    "Modèle": "ELITEBOOK - ",
    "New ID": "SL",
    "New S/N": "PT-PC-12345",
    "Old ID": "SL",
    "Old S/N": "PT-PC-67890",
    "Old Marque": "LENOVO",  # Nouveau champ
    "Old Modèle": "THINKPAD T"  # Nouveau champ
}

# Obtenez la date du jour et formatez-la sous forme "Day-Month-Year"
today_date = datetime.datetime.now().strftime("%d-%m-%Y")

# Liste pour stocker les entrées
entries = []

# Créer les champs d'entrée avec valeurs par défaut
for field in default_values:
    frame = tk.Frame(root)
    frame.pack(pady=2)
    
    label = tk.Label(frame, text=field, width=20, anchor="w")
    label.pack(side=tk.LEFT)
    
    entry = tk.Entry(frame, width=30)
    entry.pack(side=tk.RIGHT)
    
    # Si c'est le champ "Date", on insère la date du jour
    if field == "Date":
        entry.insert(0, today_date)
    else:
        entry.insert(0, default_values[field])  # Autres champs avec leurs valeurs par défaut
    
    entries.append(entry)  # Ajouter l'entry dans la liste

# Affecter les champs d'entrée à des variables
entry_nom, entry_prenom, entry_date, entry_marque, entry_modele, entry_new_id, entry_new_sn, entry_old_id, entry_old_sn, entry_old_marque, entry_old_modele = entries

# Bouton pour générer les documents
generate_button = tk.Button(root, text="Générer Documents", command=submit_form)
generate_button.pack(pady=10)

root.mainloop()

# Author: Sofiane Wattiez
# Date: 12.03.2025 
# Description: This script Automate generate 2 wordx documents based on templates and user input.
